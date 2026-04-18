import os
import io
import uuid
import time
import json
import logging
import tempfile
import threading
import traceback
from pathlib import Path
from flask import Flask, render_template, request, jsonify, send_file
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.enums import TA_LEFT

from dotenv import load_dotenv
load_dotenv()  # loads .env file if present (no-op in production where env vars are set directly)

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.StreamHandler(),
        logging.FileHandler("raydex.log")
    ]
)

app = Flask(__name__)
app.config["MAX_CONTENT_LENGTH"] = 500 * 1024 * 1024
app.config["SECRET_KEY"] = os.environ.get("SECRET_KEY", os.urandom(32))
app.config["JSON_SORT_KEYS"] = False

GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")

ALLOWED_AUDIO = {"mp3", "wav", "m4a", "ogg", "flac", "aac", "opus", "wma"}
ALLOWED_VIDEO = {"mp4", "mkv", "mov", "avi", "webm", "wmv", "flv", "m4v", "3gp"}
ALLOWED_EXTENSIONS = ALLOWED_AUDIO | ALLOWED_VIDEO

# Formats the Groq Whisper API accepts natively (no conversion needed)
SUPPORTED_FORMATS = {"mp3", "mp4", "mpeg", "mpga", "m4a", "wav", "webm", "ogg", "flac", "opus"}

# Groq Whisper API hard limit
MAX_FILE_BYTES = 25 * 1024 * 1024

jobs: dict = {}
jobs_lock = threading.Lock()

UPLOAD_TMP = Path(tempfile.gettempdir()) / "raydex_transcriber"
UPLOAD_TMP.mkdir(parents=True, exist_ok=True)


def allowed_file(filename: str) -> bool:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    return ext in ALLOWED_EXTENSIONS


def get_extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def format_timestamp(seconds: float, srt: bool = False) -> str:
    h  = int(seconds // 3600)
    m  = int((seconds % 3600) // 60)
    s  = int(seconds % 60)
    ms = int(round((seconds % 1) * 1000))
    if srt:
        return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"
    return f"{h:02d}:{m:02d}:{s:02d}"


def build_srt(segments) -> str:
    lines = []
    for i, seg in enumerate(segments, 1):
        lines.append(
            f"{i}\n{format_timestamp(seg['start'], srt=True)} --> "
            f"{format_timestamp(seg['end'], srt=True)}\n{seg['text'].strip()}\n"
        )
    return "\n".join(lines)


def build_txt(segments, include_timestamps: bool = True) -> str:
    if not include_timestamps:
        return "\n".join(seg["text"].strip() for seg in segments)
    return "\n".join(
        f"[{format_timestamp(seg['start'])}] {seg['text'].strip()}"
        for seg in segments
    )


def build_docx(segments, filename="transcript", language="", include_timestamps=True) -> io.BytesIO:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Georgia"
    style.font.size = Pt(11)

    title = doc.add_heading(level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Raydex Transcriber")
    run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x6E)
    run.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.add_run(f"File: {filename}\n")
    if language:
        meta.add_run(f"Detected Language: {language.upper()}\n")
    meta.add_run(f"Generated: {time.strftime('%Y-%m-%d %H:%M')}")
    meta.style.font.size = Pt(9)
    doc.add_paragraph()

    for seg in segments:
        p = doc.add_paragraph()
        if include_timestamps:
            ts_run = p.add_run(f"[{format_timestamp(seg['start'])}]  ")
            ts_run.font.color.rgb = RGBColor(0x1A, 0x7A, 0x6E)
            ts_run.font.size = Pt(9)
            ts_run.bold = True
        p.add_run(seg["text"].strip())

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_pdf(segments, filename="transcript", language="", include_timestamps=True) -> io.BytesIO:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=2.5*cm, rightMargin=2.5*cm,
        topMargin=2.5*cm,  bottomMargin=2.5*cm,
    )
    teal  = colors.HexColor("#1a7a6e")
    ink   = colors.HexColor("#1c1814")
    muted = colors.HexColor("#5c4f3d")

    story = [
        Paragraph("Raydex Transcriber", ParagraphStyle(
            "RaydexTitle", fontName="Helvetica-Bold", fontSize=20,
            textColor=teal, spaceAfter=6, leading=24)),
        HRFlowable(width="100%", thickness=1.5, color=teal, spaceAfter=8),
    ]

    meta_lines = [f"<b>File:</b> {filename}"]
    if language:
        meta_lines.append(f"<b>Language:</b> {language.upper()}")
    meta_lines.append(f"<b>Generated:</b> {time.strftime('%Y-%m-%d %H:%M')}")
    meta_lines.append(f"<b>Segments:</b> {len(segments)}")
    story.append(Paragraph("  &nbsp;&nbsp;".join(meta_lines), ParagraphStyle(
        "RaydexMeta", fontName="Helvetica", fontSize=9,
        textColor=muted, spaceAfter=4, leading=14)))
    story.append(Spacer(1, 0.4*cm))
    story.append(HRFlowable(width="100%", thickness=0.5,
                             color=colors.HexColor("#e8dfd0"), spaceAfter=12))

    body_style = ParagraphStyle(
        "RaydexBody", fontName="Helvetica", fontSize=10.5,
        textColor=ink, leading=16, spaceAfter=8)

    for seg in segments:
        if include_timestamps:
            txt = (f'<font name="Courier-Bold" color="#1a7a6e" size="8">'
                   f'[{format_timestamp(seg["start"])}]</font>  {seg["text"].strip()}')
        else:
            txt = seg["text"].strip()
        story.append(Paragraph(txt, body_style))

    story += [
        Spacer(1, 0.5*cm),
        HRFlowable(width="100%", thickness=0.5,
                   color=colors.HexColor("#e8dfd0"), spaceAfter=6),
        Paragraph(
            "Generated by <b>Raydex Transcriber</b> &middot; raydexhub.com &middot; Powered by Groq Whisper",
            ParagraphStyle("Footer", fontName="Helvetica", fontSize=8,
                           textColor=muted, alignment=1)),
    ]

    doc.build(story)
    buf.seek(0)
    return buf


def _call_groq_whisper(filepath: str, task: str, language: str) -> dict:
    """Send file to Groq Whisper API, return normalised result dict."""
    client = Groq(api_key=GROQ_API_KEY)

    with open(filepath, "rb") as audio_file:
        if task == "translate":
            response = client.audio.translations.create(
                file=audio_file,
                model="whisper-large-v3-turbo",
                response_format="verbose_json",
            )
        else:
            kwargs = {
                "model": "whisper-large-v3-turbo",
                "response_format": "verbose_json",
            }
            if language and language != "auto":
                kwargs["language"] = language
            response = client.audio.transcriptions.create(file=audio_file, **kwargs)

    # Log the response shape for debugging
    logging.info("[whisper] response type=%s, has segments=%s, text_len=%d",
                 type(response).__name__,
                 hasattr(response, "segments"),
                 len(getattr(response, "text", "") or ""))

    raw_segments = getattr(response, "segments", None)
    # Groq sometimes returns None, [], or a list of dicts/objects
    if not raw_segments:
        raw_segments = []
    logging.info("[whisper] raw_segments type=%s count=%d first=%s",
                 type(raw_segments[0]).__name__ if raw_segments else "–",
                 len(raw_segments),
                 repr(raw_segments[0]) if raw_segments else "–")

    def _get(seg, key, default=0):
        """Read a field from either a dict or an object (Groq returns both)."""
        if isinstance(seg, dict):
            return seg.get(key, default)
        return getattr(seg, key, default)

    segments = [
        {
            "id":    i,
            "start": round(float(_get(seg, "start", 0)), 2),
            "end":   round(float(_get(seg, "end",   0)), 2),
            "text":  str(_get(seg, "text", "")).strip(),
        }
        for i, seg in enumerate(raw_segments)
    ]

    # Detect if all timestamps are zero (Groq API returned no timing data)
    all_zero = segments and all(s["start"] == 0.0 and s["end"] == 0.0 for s in segments)
    if all_zero:
        logging.warning("[whisper] All segment timestamps are 0.0 — Groq did not return timing data. "
                        "Estimating timestamps from word count.")
        # Rough estimate: average spoken word rate ~2.5 words/second
        WORDS_PER_SEC = 2.5
        cursor = 0.0
        for seg in segments:
            word_count = max(1, len(seg["text"].split()))
            duration = word_count / WORDS_PER_SEC
            seg["start"] = round(cursor, 2)
            seg["end"]   = round(cursor + duration, 2)
            cursor += duration

    full_text = getattr(response, "text", "") or ""
    if not segments and full_text:
        segments = [{"id": 0, "start": 0.0, "end": 0.0, "text": full_text.strip()}]

    return {
        "text":     full_text,
        "language": getattr(response, "language", "") or "",
        "segments": segments,
    }


def transcribe_job(job_id: str, filepath: str, task: str, language: str):
    try:
        with jobs_lock:
            jobs[job_id]["status"]   = "transcribing"
            jobs[job_id]["progress"] = 30

        logging.info("[job:%s] Sending to Groq Whisper API (task=%s)...", job_id, task)
        result = _call_groq_whisper(filepath, task, language)
        segments = result["segments"]

        logging.info("[job:%s] Done — %d segments, language=%s",
                     job_id, len(segments), result["language"])

        with jobs_lock:
            jobs[job_id].update({
                "status":   "done",
                "progress": 100,
                "segments": segments,
                "text":     result["text"],
                "language": result["language"],
                "duration": segments[-1]["end"] if segments else 0,
            })

    except Exception as e:
        tb  = traceback.format_exc()
        raw = str(e)
        logging.error("[job:%s] Transcription failed:\n%s", job_id, tb)

        raw_lower = raw.lower()
        if "invalid_api_key" in raw_lower or "incorrect api key" in raw_lower:
            friendly = ("Invalid Groq API key. "
                        "Set the GROQ_API_KEY environment variable and restart the server.")
        elif "insufficient_quota" in raw_lower or "rate limit" in raw_lower:
            friendly = ("Groq API quota exceeded or rate-limited. "
                        "The free tier allows 7,200 minutes/day — check console.groq.com for usage.")
        elif "413" in raw or "maximum content size" in raw_lower or "file size" in raw_lower:
            friendly = ("File is too large for the Groq API (25 MB limit). "
                        "Please compress or trim the file and try again.")
        elif "connection" in raw_lower or "timeout" in raw_lower:
            friendly = ("Could not reach the Groq API. "
                        "Check the server's internet connection and try again.")
        else:
            friendly = raw

        with jobs_lock:
            jobs[job_id].update({"status": "error", "error": friendly})

    finally:
        try:
            Path(filepath).unlink(missing_ok=True)
            logging.info("[job:%s] Temp file cleaned up.", job_id)
        except Exception as cleanup_err:
            logging.warning("[job:%s] Temp file cleanup failed: %s", job_id, cleanup_err)


# ── Routes ────────────────────────────────────────────────────────────────────

@app.route("/")
def index():
    return render_template("index.html")


@app.route("/transcribe", methods=["POST"])
def transcribe():
    if not GROQ_API_KEY:
        return jsonify({"error":
            "Server is not configured with a Groq API key. "
            "Please contact the administrator."}), 500

    if "file" not in request.files:
        return jsonify({"error": "No file provided."}), 400

    file = request.files["file"]
    if not file.filename or not allowed_file(file.filename):
        return jsonify({"error": "Unsupported file type."}), 400

    ext = get_extension(file.filename)
    if ext not in SUPPORTED_FORMATS:
        return jsonify({"error":
            f".{ext} files are not supported. "
            "Please use MP3, M4A, WAV, FLAC, OGG, OPUS, or WEBM."}), 400

    # Check file size before reading it all into memory
    file.stream.seek(0, 2)
    file_size = file.stream.tell()
    file.stream.seek(0)
    if file_size > MAX_FILE_BYTES:
        return jsonify({"error":
            f"File is {file_size / 1048576:.1f} MB — the 25 MB limit applies. "
            "Please compress or trim the audio and try again."}), 400

    task     = request.form.get("task",     "transcribe")
    language = request.form.get("language", "auto")
    if task not in ("transcribe", "translate"):
        task = "transcribe"

    # Save with explicit fsync so the background thread always finds the file
    tmp_path = str(UPLOAD_TMP / f"{uuid.uuid4()}.{ext}")
    with open(tmp_path, "wb") as fout:
        file.stream.seek(0)
        while chunk := file.stream.read(1024 * 1024):
            fout.write(chunk)
        fout.flush()
        os.fsync(fout.fileno())

    logging.info("Saved '%s' (%d bytes) -> %s", file.filename, file_size, tmp_path)

    job_id = str(uuid.uuid4())
    with jobs_lock:
        jobs[job_id] = {
            "status":   "queued",
            "progress": 15,
            "filename": file.filename,
            "task":     task,
        }

    threading.Thread(
        target=transcribe_job,
        args=(job_id, tmp_path, task, language),
        daemon=True,
    ).start()

    return jsonify({"job_id": job_id})


@app.route("/job/<job_id>")
def job_status(job_id):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job:
        return jsonify({"error": "Job not found"}), 404
    return jsonify(job)


@app.route("/export/<job_id>/<fmt>")
def export(job_id, fmt):
    with jobs_lock:
        job = jobs.get(job_id)
    if not job or job.get("status") != "done":
        return jsonify({"error": "Transcript not ready"}), 400

    segments   = job["segments"]
    language   = job.get("language", "")
    filename   = job.get("filename", "transcript").rsplit(".", 1)[0]
    include_ts = request.args.get("timestamps", "true") == "true"

    if fmt == "txt":
        buf = io.BytesIO(build_txt(segments, include_ts).encode("utf-8"))
        return send_file(buf, mimetype="text/plain", as_attachment=True,
                         download_name=f"{filename}_transcript.txt")
    if fmt == "srt":
        buf = io.BytesIO(build_srt(segments).encode("utf-8"))
        return send_file(buf, mimetype="text/plain", as_attachment=True,
                         download_name=f"{filename}_transcript.srt")
    if fmt == "docx":
        return send_file(
            build_docx(segments, filename, language, include_ts),
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment=True,
            download_name=f"{filename}_transcript.docx",
        )
    if fmt == "pdf":
        return send_file(
            build_pdf(segments, filename, language, include_ts),
            mimetype="application/pdf", as_attachment=True,
            download_name=f"{filename}_transcript.pdf",
        )
    if fmt == "json":
        buf = io.BytesIO(json.dumps(
            {"filename": filename, "language": language, "segments": segments},
            indent=2, ensure_ascii=False,
        ).encode("utf-8"))
        return send_file(buf, mimetype="application/json", as_attachment=True,
                         download_name=f"{filename}_transcript.json")

    return jsonify({"error": "Unknown format"}), 400


if __name__ == "__main__":
    logging.info("Starting Raydex Transcriber...")
    logging.info("Temp upload directory: %s", UPLOAD_TMP)
    if not GROQ_API_KEY:
        logging.error(
            "GROQ_API_KEY is not set. Set it before starting:\n"
            "  Windows : set GROQ_API_KEY=gsk_...\n"
            "  Linux   : export GROQ_API_KEY=gsk_..."
        )
    else:
        logging.info("Groq API key loaded (...%s)", GROQ_API_KEY[-4:])
    app.run(debug=False, host="0.0.0.0", port=5001)
